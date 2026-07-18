"""Per-document text acquisition: get the best available text for every page,
using the PDF text layer when present and falling back to Tesseract OCR for
scanned / image-only pages. Produces per-page provenance (text vs ocr) and OCR
confidence so later stages can gate low-quality candidates for human review.

OCR is never treated as authoritative here — it only *fills* pages the text
layer could not provide, and always carries its confidence forward.
"""

from __future__ import annotations

from statistics import mean
from typing import Any

from ..extraction import image_only, ocr, pdf_render, pdf_text, vlm


def acquire_document(
    name: str, path: str, kind: str, settings, defer_vlm: bool = False
) -> dict[str, Any]:
    """defer_vlm=True runs the CPU stages only (text layer + OCR) and records
    which pages still need the VLM rescue; apply_vlm_rescue() then runs the
    deferred escalations GROUPED, so the vision model never competes with the
    extraction model on the same GPU (grouping preserves rescue reliability).

    Word (.docx) is handled by a direct text reader - it already has a clean
    text layer, so it skips OCR and VLM entirely."""
    if kind == "docx":
        doc = _acquire_docx(name, path, settings)
    elif kind == "image":
        doc = _acquire_image(name, path, settings, defer_vlm)
    else:
        doc = _acquire_pdf(name, path, settings, defer_vlm)
    doc["_path"] = path  # internal: needed by the deferred VLM phase
    return doc


def _acquire_docx(name: str, path: str, settings) -> dict[str, Any]:
    """Extract text from a Word .docx (paragraphs + table cells) into a single
    text 'page'. No OCR/VLM: docx carries a real text layer, so this is the
    fastest and highest-fidelity input. A corrupt/empty file yields empty text,
    which the normal gating flags as NEEDS_REVIEW (never a crash)."""
    text = ""
    try:
        import docx  # python-docx

        d = docx.Document(path)
        parts: list[str] = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        # CVs often put skills/experience in tables; include cell text, de-duping
        # the repeated cell objects that python-docx returns for merged cells.
        for tbl in d.tables:
            for row in tbl.rows:
                seen, cells = set(), []
                for c in row.cells:
                    t = (c.text or "").strip()
                    if t and t not in seen:
                        seen.add(t)
                        cells.append(t)
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
    except Exception:
        text = ""  # unreadable docx -> empty -> gated to NEEDS_REVIEW downstream
    page = {"page": 1, "text": text, "source": "text", "ocr_confidence": None}
    return _summarize(name, "docx", [page])


def _needs_vlm(page: dict[str, Any], settings) -> bool:
    if page["source"] != "ocr":
        return False
    conf = page.get("ocr_confidence") or 0.0
    return (
        conf < settings.ocr_min_confidence
        or len((page["text"] or "").strip()) < settings.ocr_text_min
    )


def apply_vlm_rescue(doc: dict[str, Any], settings) -> dict[str, Any]:
    """Run the VLM escalations recorded by defer_vlm acquisition. Identical
    conditions and behavior to inline escalation; only the scheduling differs.
    Refreshes the document summary counters afterwards."""
    if not vlm.available(settings) or not doc.get("_vlm_pending"):
        doc.pop("_vlm_pending", None)
        return doc
    path = doc.get("_path")
    if not path:
        doc.pop("_vlm_pending", None)
        return doc
    escalated = 0
    for p in doc["pages"]:
        if escalated >= settings.vlm_max_pages_per_doc:
            break
        if p["page"] not in doc["_vlm_pending"]:
            continue
        if doc["kind"] == "image":
            try:
                with open(path, "rb") as f:
                    png = f.read()
            except OSError:
                continue
        else:
            try:
                png = pdf_render.render_png(path, p["page"] - 1, dpi=settings.ocr_dpi)
            except Exception:
                continue
        vtext = vlm.transcribe_png(settings, png)
        escalated += 1
        if vtext:
            p["text"] = vtext
            p["source"] = "vlm"
    doc.pop("_vlm_pending", None)
    refreshed = _summarize(doc["file"], doc["kind"], doc["pages"])
    refreshed["truncated"] = doc.get("truncated", False)
    refreshed["_path"] = path
    return refreshed


def _acquire_image(name: str, path: str, settings, defer_vlm: bool = False) -> dict[str, Any]:
    text, conf = ("", 0.0)
    if settings.ocr_mode != "off":
        text, conf = ocr.ocr_image_file(path, langs=settings.ocr_langs)
    page = {"page": 1, "text": text, "source": "ocr", "ocr_confidence": round(conf, 1)}
    needs = vlm.available(settings) and _needs_vlm(page, settings)
    if needs and defer_vlm:
        doc = _summarize(name, "image", [page])
        doc["_vlm_pending"] = [1]
        return doc
    if needs:
        try:
            with open(path, "rb") as f:
                vtext = vlm.transcribe_png(settings, f.read())
        except OSError:
            vtext = ""
        if vtext:
            page = {"page": 1, "text": vtext, "source": "vlm", "ocr_confidence": round(conf, 1)}
    return _summarize(name, "image", [page])


def _acquire_pdf(name: str, path: str, settings, defer_vlm: bool = False) -> dict[str, Any]:
    text_pages = pdf_text.read_pdf(path)
    info = {d["page"]: d for d in image_only.page_image_info(path)}

    # If the text layer is entirely missing (e.g. fully scanned) drive page count
    # from the image inspector so we still OCR every page.
    if not text_pages and info:
        text_pages = [{"page": p, "text": "", "tables": []} for p in sorted(info)]

    text_pages = text_pages[: settings.max_pages_per_doc]
    truncated = len(info) > settings.max_pages_per_doc if info else False

    # Decide which pages need OCR.
    need_ocr: list[int] = []
    if settings.ocr_mode != "off":
        for tp in text_pages:
            pg = tp["page"]
            thin = len((tp.get("text") or "").strip()) < settings.ocr_text_min
            scanned = info.get(pg, {}).get("image_only", False)
            if settings.ocr_mode == "force" or thin or scanned:
                need_ocr.append(pg - 1)  # 0-based for the renderer

    ocr_res = (
        ocr.ocr_pages(path, need_ocr, langs=settings.ocr_langs, dpi=settings.ocr_dpi)
        if need_ocr
        else {}
    )

    pages: list[dict[str, Any]] = []
    for tp in text_pages:
        pg = tp["page"]
        layer_text = (tp.get("text") or "").strip()
        otext, oconf = ocr_res.get(pg - 1, ("", 0.0))
        otext = (otext or "").strip()
        # Use OCR only when the text layer was insufficient and OCR actually read something.
        if otext and (settings.ocr_mode == "force" or len(layer_text) < settings.ocr_text_min):
            pages.append(
                {"page": pg, "text": otext, "source": "ocr", "ocr_confidence": round(oconf, 1)}
            )
        elif layer_text or not need_ocr:
            pages.append({"page": pg, "text": layer_text, "source": "text", "ocr_confidence": None})
        else:
            # OCR was attempted but produced nothing usable
            pages.append(
                {"page": pg, "text": layer_text, "source": "ocr", "ocr_confidence": round(oconf, 1)}
            )

    # Escalation ladder, final rung: pages OCR could not read confidently go to
    # the LOCAL VLM transcriber (bounded per document; graceful on failure).
    # With defer_vlm the pages are only RECORDED here and rescued later in a
    # grouped phase (apply_vlm_rescue), keeping vision and extraction models
    # from competing for the GPU.
    pending = (
        [p["page"] for p in pages if _needs_vlm(p, settings)] if vlm.available(settings) else []
    )
    if pending and defer_vlm:
        doc = _summarize(name, "pdf", pages)
        doc["truncated"] = truncated
        doc["_vlm_pending"] = pending
        return doc
    if pending:
        escalated = 0
        for p in pages:
            if escalated >= settings.vlm_max_pages_per_doc:
                break
            if p["page"] not in pending:
                continue
            try:
                png = pdf_render.render_png(path, p["page"] - 1, dpi=settings.ocr_dpi)
            except Exception:
                continue
            vtext = vlm.transcribe_png(settings, png)
            escalated += 1
            if vtext:
                p["text"] = vtext
                p["source"] = "vlm"

    doc = _summarize(name, "pdf", pages)
    doc["truncated"] = truncated
    return doc


def _summarize(name: str, kind: str, pages: list[dict[str, Any]]) -> dict[str, Any]:
    # VLM-rescued pages are counted separately: their OCR confidence no longer
    # describes the text actually used (which is the VLM transcription).
    ocr_confs = [
        p["ocr_confidence"]
        for p in pages
        if p["source"] == "ocr" and p["ocr_confidence"] is not None
    ]
    vlm_pages = sum(1 for p in pages if p["source"] == "vlm")
    text_chars = sum(len(p["text"] or "") for p in pages)
    return {
        "file": name,
        "kind": kind,
        "page_count": len(pages),
        "pages": pages,
        "ocr_page_count": len(ocr_confs),
        "vlm_page_count": vlm_pages,
        "mean_ocr_confidence": round(mean(ocr_confs), 1) if ocr_confs else None,
        "text_chars": text_chars,
        "empty": text_chars == 0,
    }
